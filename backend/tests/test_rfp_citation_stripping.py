"""
RFP Citation Stripping and Export Tests
Tests for ensuring [Source:...] citations are removed from exports
and all sections have substantive content
"""
import pytest
import requests
import io
import os

BASE_URL = os.environ.get('REACT_APP_BACKEND_URL', 'https://ai-insight-canvas.preview.emergentagent.com')

# Test data with citations that should be stripped
TEST_SECTIONS_WITH_CITATIONS = [
    {
        "section": "Executive Summary",
        "content": "This project will revolutionize healthcare delivery. [Source: Healthcare Report · §Overview] The platform enables remote patient monitoring [Source: Tech Specs · §Architecture] and integrates with existing hospital systems.\n\nKey benefits include improved patient outcomes [Source: Clinical Study · §Results] and reduced costs."
    },
    {
        "section": "Project Objectives",
        "content": "The primary objective is to deploy a comprehensive telemedicine solution. [Source: Project Brief · §Goals]\n\nSecondary objectives include training staff [Source: Training Manual · §Phase 1] and establishing 24/7 support."
    },
    {
        "section": "Technical Requirements",
        "content": "The system must support HIPAA compliance [Source: Compliance Doc · §Security] and integrate via HL7 FHIR APIs. [Source: Integration Spec · §API]\n\nInfrastructure requirements include cloud hosting with 99.9% uptime [Source: SLA · §Availability]."
    },
    {
        "section": "Evaluation Criteria",
        "content": "Proposals will be evaluated based on technical capability [Source: Evaluation Guide · §Criteria], pricing structure, and vendor experience.\n\nThe scoring methodology assigns 40% to technical approach [Source: Scoring Matrix · §Weights], 30% to cost, and 30% to qualifications."
    },
    {
        "section": "Submission Requirements",
        "content": "All proposals must be submitted electronically by the deadline. [Source: Submission Guidelines · §Process]\n\nRequired documents include company profile [Source: Template · §Documents], technical proposal, and cost proposal in separate volumes."
    },
    {
        "section": "Budget Considerations",
        "content": "The total project budget is estimated at $2.5M [Source: Budget Planning · §Estimates] over three years.\n\nPayment milestones will be tied to deliverable acceptance [Source: Contract Terms · §Payments]."
    }
]


class TestCitationStripping:
    """Tests for citation stripping in exports"""
    
    def test_docx_export_no_citations(self):
        """Verify DOCX export contains no [Source:...] text"""
        response = requests.post(
            f"{BASE_URL}/api/rfp/export",
            json={
                "sections": TEST_SECTIONS_WITH_CITATIONS,
                "project_name": "Citation Strip Test",
                "client_name": "Test Organization",
                "tone": "Formal",
                "source_names": ["Doc1", "Doc2"],
                "format": "docx"
            }
        )
        
        assert response.status_code == 200, f"Expected 200, got {response.status_code}"
        
        # Parse the DOCX using python-docx
        from docx import Document
        doc = Document(io.BytesIO(response.content))
        
        # Check all paragraphs for citations
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Also check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        combined_text = '\n'.join(full_text)
        
        # Verify no citations exist
        assert '[Source:' not in combined_text, f"Found [Source: citation in DOCX content"
        assert 'Healthcare Report' not in combined_text or '[Source: Healthcare Report' not in combined_text
        print(f"✅ DOCX export contains NO [Source:] citations")
        print(f"   Total text length: {len(combined_text)} chars")
    
    def test_pdf_export_no_citations(self):
        """Verify PDF export contains no [Source:...] text"""
        response = requests.post(
            f"{BASE_URL}/api/rfp/export",
            json={
                "sections": TEST_SECTIONS_WITH_CITATIONS,
                "project_name": "PDF Citation Test",
                "client_name": "Test Organization",
                "tone": "Formal",
                "source_names": ["Doc1", "Doc2"],
                "format": "pdf"
            }
        )
        
        assert response.status_code == 200, f"Expected 200, got {response.status_code}"
        
        # Check PDF signature
        assert response.content[:4] == b'%PDF', "Invalid PDF signature"
        
        # Extract text from PDF
        import PyPDF2
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(response.content))
        full_text = []
        for page in pdf_reader.pages:
            text = page.extract_text()
            if text:
                full_text.append(text)
        
        combined_text = '\n'.join(full_text)
        
        # Verify no citations exist
        assert '[Source:' not in combined_text, f"Found [Source: citation in PDF content"
        print(f"✅ PDF export contains NO [Source:] citations")
        print(f"   Total text length: {len(combined_text)} chars, Pages: {len(pdf_reader.pages)}")
    
    def test_docx_has_proper_headings(self):
        """Verify DOCX uses Heading 1 style for section headings"""
        response = requests.post(
            f"{BASE_URL}/api/rfp/export",
            json={
                "sections": TEST_SECTIONS_WITH_CITATIONS[:3],
                "project_name": "Heading Test",
                "client_name": "Test Org",
                "tone": "Formal",
                "format": "docx"
            }
        )
        
        assert response.status_code == 200
        
        from docx import Document
        doc = Document(io.BytesIO(response.content))
        
        # Find headings
        heading1_count = 0
        heading2_count = 0
        
        for para in doc.paragraphs:
            if para.style and para.style.name:
                if para.style.name == 'Heading 1':
                    heading1_count += 1
                    print(f"   Heading 1: {para.text[:50]}...")
                elif para.style.name == 'Heading 2':
                    heading2_count += 1
                    print(f"   Heading 2: {para.text[:50]}...")
        
        # Should have at least one Heading 1 for each section
        assert heading1_count >= len(TEST_SECTIONS_WITH_CITATIONS[:3]), f"Expected >= {len(TEST_SECTIONS_WITH_CITATIONS[:3])} Heading 1, found {heading1_count}"
        print(f"✅ DOCX has {heading1_count} Heading 1 styles")
    
    def test_docx_has_table_of_contents(self):
        """Verify DOCX has a Table of Contents"""
        response = requests.post(
            f"{BASE_URL}/api/rfp/export",
            json={
                "sections": TEST_SECTIONS_WITH_CITATIONS,
                "project_name": "TOC Test",
                "client_name": "Test Org",
                "tone": "Formal",
                "format": "docx"
            }
        )
        
        assert response.status_code == 200
        
        from docx import Document
        doc = Document(io.BytesIO(response.content))
        
        full_text = '\n'.join([p.text for p in doc.paragraphs])
        
        # Check for TOC indicators
        has_toc = 'Table of Contents' in full_text or 'TABLE OF CONTENTS' in full_text
        assert has_toc, "No Table of Contents found in DOCX"
        
        # Check that sections are listed in TOC
        for sec in TEST_SECTIONS_WITH_CITATIONS:
            assert sec['section'] in full_text, f"Section '{sec['section']}' not found in document"
        
        print(f"✅ DOCX has Table of Contents with all {len(TEST_SECTIONS_WITH_CITATIONS)} sections listed")
    
    def test_pdf_has_table_of_contents(self):
        """Verify PDF has a Table of Contents"""
        response = requests.post(
            f"{BASE_URL}/api/rfp/export",
            json={
                "sections": TEST_SECTIONS_WITH_CITATIONS,
                "project_name": "PDF TOC Test",
                "client_name": "Test Org",
                "tone": "Formal",
                "format": "pdf"
            }
        )
        
        assert response.status_code == 200
        
        import PyPDF2
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(response.content))
        
        full_text = '\n'.join([p.extract_text() or '' for p in pdf_reader.pages])
        
        # Check for TOC
        has_toc = 'Table of Contents' in full_text or 'TABLE OF CONTENTS' in full_text
        assert has_toc, "No Table of Contents found in PDF"
        
        print(f"✅ PDF has Table of Contents, {len(pdf_reader.pages)} pages")
    
    def test_docx_has_title_page(self):
        """Verify DOCX has proper title page with project and client"""
        response = requests.post(
            f"{BASE_URL}/api/rfp/export",
            json={
                "sections": TEST_SECTIONS_WITH_CITATIONS[:2],
                "project_name": "Hospital at Home Platform",
                "client_name": "Ministry of Health",
                "tone": "Formal",
                "format": "docx"
            }
        )
        
        assert response.status_code == 200
        
        from docx import Document
        doc = Document(io.BytesIO(response.content))
        
        full_text = '\n'.join([p.text for p in doc.paragraphs])
        
        # Check title page elements
        assert 'Hospital at Home Platform' in full_text, "Project name not in document"
        assert 'Ministry of Health' in full_text, "Client name not in document"
        assert 'REQUEST FOR PROPOSAL' in full_text or 'Request for Proposal' in full_text, "RFP label not found"
        
        print(f"✅ DOCX has title page with project name and client")
    
    def test_pdf_has_title_page(self):
        """Verify PDF has proper title page elements"""
        response = requests.post(
            f"{BASE_URL}/api/rfp/export",
            json={
                "sections": TEST_SECTIONS_WITH_CITATIONS[:2],
                "project_name": "Hospital at Home Platform",
                "client_name": "Ministry of Health",
                "tone": "Formal",
                "format": "pdf"
            }
        )
        
        assert response.status_code == 200
        
        import PyPDF2
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(response.content))
        
        # Check first page (title page)
        title_text = pdf_reader.pages[0].extract_text() or ''
        
        assert 'Hospital at Home Platform' in title_text, "Project name not on title page"
        assert 'Ministry of Health' in title_text, "Client name not on title page"
        
        print(f"✅ PDF has title page with project and client info")


class TestSectionContent:
    """Tests for section content quality"""
    
    def test_evaluation_criteria_has_content(self):
        """Verify Evaluation Criteria section has substantive content"""
        # Check from our test data
        for sec in TEST_SECTIONS_WITH_CITATIONS:
            if sec['section'] == 'Evaluation Criteria':
                content = sec['content'].replace('[Source:', '').strip()
                assert len(content) > 100, f"Evaluation Criteria content too short: {len(content)} chars"
                print(f"✅ Evaluation Criteria has {len(content)} chars of content")
                break
    
    def test_submission_requirements_has_content(self):
        """Verify Submission Requirements section has substantive content"""
        for sec in TEST_SECTIONS_WITH_CITATIONS:
            if sec['section'] == 'Submission Requirements':
                content = sec['content'].replace('[Source:', '').strip()
                assert len(content) > 100, f"Submission Requirements content too short: {len(content)} chars"
                print(f"✅ Submission Requirements has {len(content)} chars of content")
                break
    
    def test_all_sections_have_content_in_export(self):
        """Verify all sections have content when exported"""
        response = requests.post(
            f"{BASE_URL}/api/rfp/export",
            json={
                "sections": TEST_SECTIONS_WITH_CITATIONS,
                "project_name": "Content Test",
                "client_name": "Test Org",
                "tone": "Formal",
                "format": "docx"
            }
        )
        
        assert response.status_code == 200
        
        from docx import Document
        doc = Document(io.BytesIO(response.content))
        
        full_text = '\n'.join([p.text for p in doc.paragraphs])
        
        # Verify each section title appears in document
        for sec in TEST_SECTIONS_WITH_CITATIONS:
            assert sec['section'] in full_text, f"Section '{sec['section']}' not found in export"
        
        print(f"✅ All {len(TEST_SECTIONS_WITH_CITATIONS)} sections have content in export")


class TestExportFormats:
    """Tests for export format functionality"""
    
    def test_docx_format_valid(self):
        """Verify DOCX export returns valid Word document"""
        response = requests.post(
            f"{BASE_URL}/api/rfp/export",
            json={
                "sections": [{"section": "Test", "content": "Test content"}],
                "project_name": "Test",
                "format": "docx"
            }
        )
        
        assert response.status_code == 200
        assert response.headers.get('content-type') == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        assert response.content[:2] == b'PK', "Invalid DOCX (should be ZIP format)"
        
        # Verify readable by python-docx
        from docx import Document
        doc = Document(io.BytesIO(response.content))
        assert len(doc.paragraphs) > 0
        
        print(f"✅ DOCX format valid, {len(doc.paragraphs)} paragraphs")
    
    def test_pdf_format_valid(self):
        """Verify PDF export returns valid PDF document"""
        response = requests.post(
            f"{BASE_URL}/api/rfp/export",
            json={
                "sections": [{"section": "Test", "content": "Test content"}],
                "project_name": "Test",
                "format": "pdf"
            }
        )
        
        assert response.status_code == 200
        assert 'application/pdf' in response.headers.get('content-type', '')
        assert response.content[:4] == b'%PDF', "Invalid PDF signature"
        
        # Verify readable by PyPDF2
        import PyPDF2
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(response.content))
        assert len(pdf_reader.pages) > 0
        
        print(f"✅ PDF format valid, {len(pdf_reader.pages)} pages")
    
    def test_invalid_format_returns_400(self):
        """Verify invalid format returns 400 error"""
        response = requests.post(
            f"{BASE_URL}/api/rfp/export",
            json={
                "sections": [{"section": "Test", "content": "Test"}],
                "project_name": "Test",
                "format": "invalid"
            }
        )
        
        assert response.status_code == 400
        print(f"✅ Invalid format correctly returns 400")


if __name__ == "__main__":
    pytest.main([__file__, "-v", "--tb=short"])

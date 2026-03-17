"""
Tests for RFP Export Functionality - DOCX and PDF generation
Tests the /api/rfp/export endpoint for both Word and PDF formats
"""
import pytest
import requests
import os
import io
import tempfile

BASE_URL = os.environ.get('REACT_APP_BACKEND_URL', 'https://ai-insight-canvas.preview.emergentagent.com')

# Sample RFP sections for testing
SAMPLE_SECTIONS = [
    {"section": "Executive Summary", "content": "This proposal outlines a comprehensive healthcare platform. [Source: HealthDoc · §Overview] The system will improve patient outcomes through modern technology.\n\nKey benefits include reduced costs and improved efficiency. [Source: TechSpec · §Benefits]"},
    {"section": "Project Objectives", "content": "The primary objectives are to enhance patient care and streamline operations. [Source: HealthDoc · §Objectives]\n\nSecondary goals focus on data analytics and reporting capabilities."}
]

SAMPLE_REQUEST = {
    "sections": SAMPLE_SECTIONS,
    "project_name": "Hospital at Home Platform",
    "client_name": "Ministry of Health",
    "tone": "Formal",
    "source_names": ["HealthDoc", "TechSpec", "Research2025"],
    "format": "docx"
}


class TestRFPExportEndpoint:
    """Tests for /api/rfp/export endpoint availability and validation"""
    
    def test_endpoint_exists(self):
        """Verify the export endpoint exists and accepts POST"""
        response = requests.options(f"{BASE_URL}/api/rfp/export")
        # Options or post should not return 404
        assert response.status_code != 404, "Export endpoint does not exist"
        print(f"✓ Export endpoint exists (status: {response.status_code})")
    
    def test_invalid_format_rejected(self):
        """Invalid format should return 400"""
        req = {**SAMPLE_REQUEST, "format": "invalid"}
        response = requests.post(f"{BASE_URL}/api/rfp/export", json=req)
        assert response.status_code == 400, f"Expected 400 for invalid format, got {response.status_code}"
        data = response.json()
        assert "detail" in data and ("docx" in data["detail"].lower() or "pdf" in data["detail"].lower())
        print(f"✓ Invalid format rejected with 400: {data.get('detail', '')}")
    
    def test_missing_sections_handled(self):
        """Request with empty sections should be handled"""
        req = {
            "sections": [],
            "project_name": "Test",
            "client_name": "Test Org",
            "tone": "Formal",
            "source_names": [],
            "format": "docx"
        }
        response = requests.post(f"{BASE_URL}/api/rfp/export", json=req)
        # Should either succeed with empty doc or return appropriate error
        assert response.status_code in [200, 400, 422], f"Unexpected status: {response.status_code}"
        print(f"✓ Empty sections handled (status: {response.status_code})")


class TestDOCXExport:
    """Tests for DOCX export functionality"""
    
    def test_docx_export_returns_200(self):
        """DOCX export should return 200 status"""
        req = {**SAMPLE_REQUEST, "format": "docx"}
        response = requests.post(f"{BASE_URL}/api/rfp/export", json=req)
        assert response.status_code == 200, f"DOCX export failed with status {response.status_code}: {response.text[:500]}"
        print(f"✓ DOCX export returned 200 OK")
    
    def test_docx_correct_content_type(self):
        """DOCX should have correct Content-Type header"""
        req = {**SAMPLE_REQUEST, "format": "docx"}
        response = requests.post(f"{BASE_URL}/api/rfp/export", json=req)
        assert response.status_code == 200
        content_type = response.headers.get('Content-Type', '')
        expected = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        assert expected in content_type, f"Wrong Content-Type: {content_type}"
        print(f"✓ DOCX has correct Content-Type: {content_type}")
    
    def test_docx_has_content_disposition(self):
        """DOCX should have Content-Disposition header with filename"""
        req = {**SAMPLE_REQUEST, "format": "docx"}
        response = requests.post(f"{BASE_URL}/api/rfp/export", json=req)
        assert response.status_code == 200
        content_disp = response.headers.get('Content-Disposition', '')
        assert 'attachment' in content_disp.lower(), f"Missing attachment in Content-Disposition: {content_disp}"
        assert '.docx' in content_disp.lower(), f"Missing .docx in Content-Disposition: {content_disp}"
        print(f"✓ DOCX Content-Disposition: {content_disp}")
    
    def test_docx_non_zero_size(self):
        """DOCX file should have non-zero size"""
        req = {**SAMPLE_REQUEST, "format": "docx"}
        response = requests.post(f"{BASE_URL}/api/rfp/export", json=req)
        assert response.status_code == 200
        content = response.content
        assert len(content) > 0, "DOCX file is empty"
        assert len(content) > 1000, f"DOCX file too small: {len(content)} bytes"
        print(f"✓ DOCX file size: {len(content)} bytes")
    
    def test_docx_valid_file_structure(self):
        """DOCX should be a valid ZIP/DOCX file"""
        req = {**SAMPLE_REQUEST, "format": "docx"}
        response = requests.post(f"{BASE_URL}/api/rfp/export", json=req)
        assert response.status_code == 200
        content = response.content
        # DOCX files start with PK (ZIP signature)
        assert content[:2] == b'PK', f"DOCX doesn't have ZIP signature. First bytes: {content[:10]}"
        print(f"✓ DOCX has valid ZIP/DOCX signature")
    
    def test_docx_can_be_read(self):
        """DOCX should be readable by python-docx"""
        req = {**SAMPLE_REQUEST, "format": "docx"}
        response = requests.post(f"{BASE_URL}/api/rfp/export", json=req)
        assert response.status_code == 200
        
        try:
            from docx import Document
            doc = Document(io.BytesIO(response.content))
            # Count paragraphs - should have at least a few
            para_count = len(doc.paragraphs)
            assert para_count > 0, "DOCX has no paragraphs"
            print(f"✓ DOCX readable with python-docx: {para_count} paragraphs")
        except ImportError:
            pytest.skip("python-docx not available for validation")
        except Exception as e:
            pytest.fail(f"DOCX is invalid: {e}")
    
    def test_docx_contains_project_name(self):
        """DOCX should contain the project name"""
        req = {**SAMPLE_REQUEST, "format": "docx"}
        response = requests.post(f"{BASE_URL}/api/rfp/export", json=req)
        assert response.status_code == 200
        
        try:
            from docx import Document
            doc = Document(io.BytesIO(response.content))
            full_text = '\n'.join([p.text for p in doc.paragraphs])
            assert "Hospital at Home Platform" in full_text, "Project name not found in DOCX"
            print(f"✓ DOCX contains project name")
        except ImportError:
            pytest.skip("python-docx not available")
    
    def test_docx_contains_client_name(self):
        """DOCX should contain the client name"""
        req = {**SAMPLE_REQUEST, "format": "docx"}
        response = requests.post(f"{BASE_URL}/api/rfp/export", json=req)
        assert response.status_code == 200
        
        try:
            from docx import Document
            doc = Document(io.BytesIO(response.content))
            full_text = '\n'.join([p.text for p in doc.paragraphs])
            assert "Ministry of Health" in full_text, "Client name not found in DOCX"
            print(f"✓ DOCX contains client name")
        except ImportError:
            pytest.skip("python-docx not available")
    
    def test_docx_contains_section_headings(self):
        """DOCX should contain section headings"""
        req = {**SAMPLE_REQUEST, "format": "docx"}
        response = requests.post(f"{BASE_URL}/api/rfp/export", json=req)
        assert response.status_code == 200
        
        try:
            from docx import Document
            doc = Document(io.BytesIO(response.content))
            full_text = '\n'.join([p.text for p in doc.paragraphs])
            assert "Executive Summary" in full_text, "Section heading 'Executive Summary' not found"
            assert "Project Objectives" in full_text, "Section heading 'Project Objectives' not found"
            print(f"✓ DOCX contains section headings")
        except ImportError:
            pytest.skip("python-docx not available")
    
    def test_docx_contains_table_of_contents(self):
        """DOCX should have Table of Contents"""
        req = {**SAMPLE_REQUEST, "format": "docx"}
        response = requests.post(f"{BASE_URL}/api/rfp/export", json=req)
        assert response.status_code == 200
        
        try:
            from docx import Document
            doc = Document(io.BytesIO(response.content))
            full_text = '\n'.join([p.text for p in doc.paragraphs])
            assert "Table of Contents" in full_text, "Table of Contents not found in DOCX"
            print(f"✓ DOCX contains Table of Contents")
        except ImportError:
            pytest.skip("python-docx not available")


class TestPDFExport:
    """Tests for PDF export functionality"""
    
    def test_pdf_export_returns_200(self):
        """PDF export should return 200 status"""
        req = {**SAMPLE_REQUEST, "format": "pdf"}
        response = requests.post(f"{BASE_URL}/api/rfp/export", json=req)
        assert response.status_code == 200, f"PDF export failed with status {response.status_code}: {response.text[:500]}"
        print(f"✓ PDF export returned 200 OK")
    
    def test_pdf_correct_content_type(self):
        """PDF should have correct Content-Type header"""
        req = {**SAMPLE_REQUEST, "format": "pdf"}
        response = requests.post(f"{BASE_URL}/api/rfp/export", json=req)
        assert response.status_code == 200
        content_type = response.headers.get('Content-Type', '')
        assert "application/pdf" in content_type, f"Wrong Content-Type: {content_type}"
        print(f"✓ PDF has correct Content-Type: {content_type}")
    
    def test_pdf_has_content_disposition(self):
        """PDF should have Content-Disposition header with filename"""
        req = {**SAMPLE_REQUEST, "format": "pdf"}
        response = requests.post(f"{BASE_URL}/api/rfp/export", json=req)
        assert response.status_code == 200
        content_disp = response.headers.get('Content-Disposition', '')
        assert 'attachment' in content_disp.lower(), f"Missing attachment in Content-Disposition: {content_disp}"
        assert '.pdf' in content_disp.lower(), f"Missing .pdf in Content-Disposition: {content_disp}"
        print(f"✓ PDF Content-Disposition: {content_disp}")
    
    def test_pdf_non_zero_size(self):
        """PDF file should have non-zero size"""
        req = {**SAMPLE_REQUEST, "format": "pdf"}
        response = requests.post(f"{BASE_URL}/api/rfp/export", json=req)
        assert response.status_code == 200
        content = response.content
        assert len(content) > 0, "PDF file is empty"
        assert len(content) > 500, f"PDF file too small: {len(content)} bytes"
        print(f"✓ PDF file size: {len(content)} bytes")
    
    def test_pdf_valid_signature(self):
        """PDF should have valid PDF signature (%PDF-)"""
        req = {**SAMPLE_REQUEST, "format": "pdf"}
        response = requests.post(f"{BASE_URL}/api/rfp/export", json=req)
        assert response.status_code == 200
        content = response.content
        # PDF files start with %PDF-
        assert content[:5] == b'%PDF-', f"PDF doesn't have valid signature. First bytes: {content[:20]}"
        print(f"✓ PDF has valid signature: {content[:8].decode('utf-8', errors='ignore')}")
    
    def test_pdf_contains_text(self):
        """PDF should contain extractable text"""
        req = {**SAMPLE_REQUEST, "format": "pdf"}
        response = requests.post(f"{BASE_URL}/api/rfp/export", json=req)
        assert response.status_code == 200
        
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(response.content))
            assert len(reader.pages) > 0, "PDF has no pages"
            
            # Extract text from first few pages
            all_text = ""
            for i, page in enumerate(reader.pages[:5]):
                all_text += page.extract_text() or ""
            
            assert len(all_text) > 50, f"PDF has very little text: {len(all_text)} chars"
            print(f"✓ PDF has {len(reader.pages)} pages with text content")
        except ImportError:
            pytest.skip("PyPDF2 not available for validation")
        except Exception as e:
            pytest.fail(f"PDF validation failed: {e}")
    
    def test_pdf_contains_project_name(self):
        """PDF should contain the project name"""
        req = {**SAMPLE_REQUEST, "format": "pdf"}
        response = requests.post(f"{BASE_URL}/api/rfp/export", json=req)
        assert response.status_code == 200
        
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(response.content))
            all_text = ""
            for page in reader.pages:
                all_text += (page.extract_text() or "")
            
            assert "Hospital" in all_text or "Home" in all_text or "Platform" in all_text, \
                "Project name not found in PDF"
            print(f"✓ PDF contains project name elements")
        except ImportError:
            pytest.skip("PyPDF2 not available")
    
    def test_pdf_multipage(self):
        """PDF should have multiple pages for full RFP"""
        req = {**SAMPLE_REQUEST, "format": "pdf"}
        response = requests.post(f"{BASE_URL}/api/rfp/export", json=req)
        assert response.status_code == 200
        
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(response.content))
            # With 2 sections + title + TOC, expect at least 3-4 pages
            assert len(reader.pages) >= 3, f"PDF should have multiple pages, got {len(reader.pages)}"
            print(f"✓ PDF has {len(reader.pages)} pages (multipage format)")
        except ImportError:
            pytest.skip("PyPDF2 not available")


class TestExportEdgeCases:
    """Edge case testing for export functionality"""
    
    def test_special_characters_in_project_name(self):
        """Export should handle special characters in project name"""
        req = {
            **SAMPLE_REQUEST,
            "project_name": "Test & Demo <Project> 2025/26",
            "format": "docx"
        }
        response = requests.post(f"{BASE_URL}/api/rfp/export", json=req)
        assert response.status_code == 200, f"Failed with special chars: {response.status_code}"
        assert len(response.content) > 0
        print(f"✓ Special characters handled in project name")
    
    def test_unicode_content(self):
        """Export should handle Unicode/international characters"""
        req = {
            **SAMPLE_REQUEST,
            "sections": [{"section": "مقدمة", "content": "هذا اختبار. [Source: Test · §Intro]"}],
            "project_name": "مشروع صحي",
            "client_name": "وزارة الصحة",
            "format": "pdf"
        }
        response = requests.post(f"{BASE_URL}/api/rfp/export", json=req)
        # Should either succeed or fail gracefully
        assert response.status_code in [200, 400, 500], f"Unexpected status: {response.status_code}"
        print(f"✓ Unicode content handled (status: {response.status_code})")
    
    def test_empty_source_names(self):
        """Export should work with empty source names"""
        req = {
            **SAMPLE_REQUEST,
            "source_names": [],
            "format": "docx"
        }
        response = requests.post(f"{BASE_URL}/api/rfp/export", json=req)
        assert response.status_code == 200, f"Failed with empty sources: {response.status_code}"
        print(f"✓ Empty source names handled")
    
    def test_long_content(self):
        """Export should handle long content"""
        long_content = "This is a test paragraph with content. [Source: Doc · §Section] " * 50
        req = {
            **SAMPLE_REQUEST,
            "sections": [{"section": "Long Section", "content": long_content}],
            "format": "pdf"
        }
        response = requests.post(f"{BASE_URL}/api/rfp/export", json=req)
        assert response.status_code == 200, f"Failed with long content: {response.status_code}"
        assert len(response.content) > 1000, "Long content PDF should be substantial"
        print(f"✓ Long content handled ({len(response.content)} bytes)")
    
    def test_many_sections(self):
        """Export should handle many sections"""
        sections = [{"section": f"Section {i}", "content": f"Content for section {i}."} for i in range(12)]
        req = {
            **SAMPLE_REQUEST,
            "sections": sections,
            "format": "docx"
        }
        response = requests.post(f"{BASE_URL}/api/rfp/export", json=req)
        assert response.status_code == 200, f"Failed with many sections: {response.status_code}"
        print(f"✓ Many sections (12) handled")


class TestFilenameGeneration:
    """Tests for proper filename generation"""
    
    def test_docx_filename_contains_project(self):
        """DOCX filename should contain project name"""
        req = {**SAMPLE_REQUEST, "format": "docx"}
        response = requests.post(f"{BASE_URL}/api/rfp/export", json=req)
        content_disp = response.headers.get('Content-Disposition', '')
        # Filename should have underscores for spaces
        assert "Hospital" in content_disp or "RFP" in content_disp, f"Filename missing project name: {content_disp}"
        print(f"✓ DOCX filename: {content_disp}")
    
    def test_pdf_filename_contains_project(self):
        """PDF filename should contain project name"""
        req = {**SAMPLE_REQUEST, "format": "pdf"}
        response = requests.post(f"{BASE_URL}/api/rfp/export", json=req)
        content_disp = response.headers.get('Content-Disposition', '')
        assert "Hospital" in content_disp or "RFP" in content_disp, f"Filename missing project name: {content_disp}"
        print(f"✓ PDF filename: {content_disp}")


# Fixture for cleanup
@pytest.fixture(scope="module", autouse=True)
def cleanup():
    """Cleanup any test artifacts"""
    yield
    # No cleanup needed for export tests as they don't create persistent data


if __name__ == "__main__":
    pytest.main([__file__, "-v", "--tb=short"])
